"""
DDQ Document Scanner - Robust Question Extraction
Handles merged cells, complex tables, content controls, and maintains
document order across PDF and Word questionnaires.
"""

import re
import os
import pandas as pd
from typing import List, Dict, Set
from datetime import datetime
import pdfplumber
from docx import Document

import config
from answer_cleaner import AnswerCleaner


class DDQScanner:
    """Robust scanner with comprehensive question extraction."""

    def __init__(self):
        self.cleaner = AnswerCleaner()
        self.seen_questions: Set[str] = set()

    def extract_from_docx(self, file_path: str) -> List[Dict]:
        """Extract questions from a Word document, including form content controls."""
        doc = Document(file_path)
        questions = []

        # First, extract from content controls (structured document tags),
        # commonly used in form templates.
        from docx.oxml.ns import qn

        for element in doc.element.body:
            sdts = element.findall('.//' + qn('w:sdt'))
            for sdt in sdts:
                text_elements = sdt.findall('.//' + qn('w:t'))
                if text_elements:
                    text = ''.join([t.text for t in text_elements if t.text])
                    if text and text.strip():
                        q_list = self._extract_questions_from_text(text)
                        for q in q_list:
                            q['source'] = 'content_control'
                            questions.append(q)

        # Extract from paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                q_list = self._extract_questions_from_text(text)
                for q in q_list:
                    q['source'] = 'paragraph'
                    questions.append(q)

        # Extract from tables - handle merged cells by tracking element ids
        for table_idx, table in enumerate(doc.tables):
            seen_cells = set()
            for row in table.rows:
                for cell in row.cells:
                    cell_id = id(cell._element)
                    if cell_id not in seen_cells:
                        seen_cells.add(cell_id)
                        cell_text = ' '.join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                        if cell_text:
                            q_list = self._extract_questions_from_text(cell_text)
                            for q in q_list:
                                q['source'] = f"table_{table_idx}"
                                questions.append(q)

        for i, q in enumerate(questions, 1):
            q['number'] = i

        return questions

    def _extract_questions_from_text(self, text: str) -> List[Dict]:
        """Extract questions from a single text block, keeping multi-part questions together."""
        questions = []

        text = self._fix_encoding(text)

        # Don't split on question marks - keeps "What is X? Please provide Y."
        # together as one logical question.
        q_data = self._process_potential_question(text)
        if q_data:
            questions.append(q_data)

        return questions

    def _fix_encoding(self, text: str) -> str:
        """Fix common character encoding issues (smart quotes, dashes, etc.)."""
        replacements = {
            '\u2018': "'",     # Left single quote
            '\u2019': "'",     # Right single quote
            '\u201c': '"',     # Left double quote
            '\u201d': '"',     # Right double quote
            '\u2013': '-',     # En dash
            '\u2014': '--',    # Em dash
        }

        for old, new in replacements.items():
            text = text.replace(old, new)

        text = ' '.join(text.split())

        return text

    def _process_potential_question(self, text: str) -> Dict:
        """Process a potential question and validate it."""
        text = text.strip()

        # Allow shorter text if it ends with a colon (form field labels)
        min_length = 10 if text.endswith(':') else 20

        if len(text) < min_length or len(text) > 2000:
            return None

        if self._is_placeholder_text(text):
            return None

        if self._is_section_header(text):
            return None

        if self._is_answer_option(text):
            return None

        text_normalized = ' '.join(text.lower().split())
        if text_normalized in self.seen_questions:
            return None

        cleaned = self._clean_question_text(text)

        if not self._is_valid_question(cleaned):
            return None

        self.seen_questions.add(text_normalized)

        return {
            'question': cleaned,
            'original_text': text
        }

    def _is_placeholder_text(self, text: str) -> bool:
        """Check if text is placeholder/template text."""
        text_lower = text.lower()

        placeholders = [
            'click or tap here',
            'click here to enter',
            'enter text here',
            'type here',
            'insert text',
            '[placeholder]',
            '[enter',
            'click to add',
        ]

        return any(placeholder in text_lower for placeholder in placeholders)

    def _is_section_header(self, text: str) -> bool:
        """Check if text is a section header or label rather than a question."""
        text_lower = text.lower()

        section_headers_strict = [
            'why this ddq is important',
            'document request list',
            'definitions',
            'instructions',
            'has the meaning used above',
            'have the meanings used above',
            'in addition to answering the questions below',
            'the following definitions apply',
            'as a general rule',
        ]

        for header in section_headers_strict:
            if text_lower.startswith(header) or (len(text) < 100 and header in text_lower):
                return True

        if len(text) < 50 and 'appendix' in text_lower:
            return True
        if text_lower.startswith('appendix'):
            return True

        if len(text) < 100 and not any(word in text_lower for word in ['please provide', 'please describe', 'please list', 'do you', 'does your', 'have you', 'has your', 'are you', 'is your']):
            header_patterns = [
                'information', 'section', 'appendix', 'part', 'chapter',
                'overview', 'summary', 'introduction', 'background', 'details',
                'controls', 'policies', 'procedures', 'compliance', 'legal',
                'regulatory', 'privacy', 'security', 'data', 'company',
                'product', 'service',
            ]

            words = text_lower.split()
            if len(words) <= 8:
                header_word_count = sum(1 for word in words if any(pattern in word for pattern in header_patterns))
                if header_word_count >= len(words) * 0.5:
                    return True

        return False

    def _is_answer_option(self, text: str) -> bool:
        """Check if text is just answer options or lists rather than a question."""
        text_lower = text.lower().strip()

        if len(text) < 60:
            simple_patterns = [
                r'^yes\s*no$',
                r'^yes\s*,?\s*and',
                r'^no\s*,?\s*and',
                r'^\s*yes\s*$',
                r'^\s*no\s*$',
                r'^yes\s+no\s+not applicable$',
            ]
            for pattern in simple_patterns:
                if re.match(pattern, text_lower):
                    return True

        # A list of options separated by several semicolons
        if len(text) > 50:
            semicolon_count = text.count(';')
            if semicolon_count >= 3:
                question_words = ['what', 'when', 'where', 'who', 'why', 'how', 'does', 'do', 'is', 'are', 'can', 'please', 'describe', 'provide', 'list', 'explain']
                if not any(text_lower.startswith(word) for word in question_words):
                    return True

        option_words = ['yes', 'no', 'not applicable', 'n/a', 'approximate number']
        word_count = len(text.split())
        option_count = sum(1 for word in option_words if word in text_lower)

        if word_count < 15 and option_count >= 2:
            return True

        return False

    def _clean_question_text(self, text: str) -> str:
        """Clean up question text (remove numbering, trailing options, etc.)."""
        text = re.sub(r'^[Q\d]+[\.\):\s]+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'^\d+[\.\)]\s+', '', text)

        text = ' '.join(text.split())

        text = re.sub(r'\s+(Yes|No|N/A|Not applicable)(\s+(Yes|No|N/A|Not applicable))+\s*$', '', text, flags=re.IGNORECASE)

        if not text.endswith(('?', '.', ':')):
            text_lower = text.lower()
            question_starters = ['what', 'when', 'where', 'who', 'why', 'how', 'does', 'do', 'is', 'are', 'can', 'could', 'would', 'should', 'will', 'have', 'has']
            if any(text_lower.startswith(word) for word in question_starters):
                text += '?'

        return text.strip()

    def _is_valid_question(self, text: str) -> bool:
        """Check if text is a valid question."""
        text_lower = text.lower()

        if not re.search(r'[a-zA-Z]{3,}', text):
            return False

        if text.isupper() and len(text) > 30:
            return False

        if text_lower.startswith('if so,') or text_lower.startswith('if yes,'):
            if '?' not in text:
                return False

        if text.endswith(':'):
            if not self._is_section_header(text):
                return True

        question_indicators = [
            '?', 'please provide', 'please describe', 'please list',
            'please explain', 'please identify', 'please indicate',
            'please attach', 'please share', 'please detail', 'please specify',
            'do you', 'does your', 'does the', 'can you', 'have you',
            'has your', 'has the', 'are you', 'is your', 'is the',
            'will you', 'would you', 'could you', 'should you', 'did you',
            'were you', 'was your', 'how many', 'how much', 'how do', 'how does',
        ]

        has_indicator = any(indicator in text_lower for indicator in question_indicators)

        question_starters = [
            'what', 'when', 'where', 'who', 'why', 'how', 'which',
            'describe', 'explain', 'list', 'provide', 'identify', 'specify'
        ]
        starts_with_question = any(text_lower.startswith(word) for word in question_starters)

        return has_indicator or starts_with_question

    def extract_from_pdf(self, file_path: str) -> List[Dict]:
        """Extract questions from a PDF (tables and free text)."""
        questions = []

        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        for row in table:
                            for cell in row:
                                if cell:
                                    q_list = self._extract_questions_from_text(str(cell))
                                    for q in q_list:
                                        q['page'] = page_num
                                        questions.append(q)

                text = page.extract_text()
                if text:
                    lines = text.split('\n')
                    for line in lines:
                        q_list = self._extract_questions_from_text(line)
                        for q in q_list:
                            q['page'] = page_num
                            questions.append(q)

        for i, q in enumerate(questions, 1):
            q['number'] = i

        return questions

    def process_document(self, file_path: str, ask_function) -> List[Dict]:
        """Process a document and answer all extracted questions."""
        print(f"\n{'='*80}")
        print(f"Processing: {file_path}")
        print(f"{'='*80}")

        # Reset for new document
        self.seen_questions = set()

        if file_path.lower().endswith('.pdf'):
            questions = self.extract_from_pdf(file_path)
        elif file_path.lower().endswith(('.docx', '.doc')):
            questions = self.extract_from_docx(file_path)
        else:
            print("Unsupported file format")
            return []

        print(f"\nExtracted {len(questions)} questions")

        if not questions:
            print("No questions found")
            return []

        print("\nFirst 10 questions:")
        for q in questions[:10]:
            print(f"  {q['number']}. {q['question'][:100]}...")

        print(f"\nAnswering {len(questions)} questions...")
        results = []

        for i, q_data in enumerate(questions, 1):
            try:
                if i % 10 == 0 or i == 1:
                    print(f"  Progress: {i}/{len(questions)}")

                answer = ask_function(q_data['question'])
                summary_answer = self.cleaner.extract_summary(answer)
                clean_answer = self.cleaner.clean_for_excel(summary_answer)
                confidence = self.cleaner.extract_confidence(answer)

                results.append({
                    'number': q_data['number'],
                    'question': q_data['question'],
                    'answer': clean_answer,
                    'confidence': confidence,
                    'page': q_data.get('page', ''),
                    'source': q_data.get('source', ''),
                    'status': 'success'
                })
            except Exception as e:
                print(f"  Error on Q{q_data['number']}: {str(e)}")
                results.append({
                    'number': q_data['number'],
                    'question': q_data['question'],
                    'answer': f"Error: {str(e)}",
                    'confidence': 'N/A',
                    'page': q_data.get('page', ''),
                    'source': q_data.get('source', ''),
                    'status': 'error'
                })

        success_count = len([r for r in results if r['status'] == 'success'])
        print(f"\nCompleted: {success_count}/{len(results)} successful")

        return results

    def export_to_csv(self, results: List[Dict], output_path: str):
        """Export results to CSV."""
        df = pd.DataFrame(results)
        columns = ['number', 'question', 'answer', 'confidence', 'page', 'source', 'status']
        df = df[[col for col in columns if col in df.columns]]
        df.to_csv(output_path, index=False, encoding='utf-8-sig')
        print(f"\nExported to: {output_path}")
        return output_path

    def process_document_from_bytes(self, file_bytes: bytes, filename: str, ask_function) -> List[Dict]:
        """
        Process a DDQ document from raw bytes and answer all questions.
        Used by the Gradio file-upload interface.
        """
        import tempfile

        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name

        try:
            results = self.process_document(tmp_path, ask_function)
            return results
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

    def format_results_as_text(self, results: List[Dict]) -> str:
        """Format results as readable text for display."""
        output = []
        output.append("=" * 80)
        output.append("DDQ AUTOMATED RESPONSE")
        output.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        output.append(f"Total Questions: {len(results)}")
        output.append("=" * 80)
        output.append("")

        for result in results:
            output.append(f"Q{result['number']}: {result['question']}")
            output.append("")
            output.append("Answer:")
            output.append(result['answer'])
            output.append("")
            if 'confidence' in result:
                output.append(f"Confidence: {result['confidence']}")
            output.append("-" * 80)
            output.append("")

        return "\n".join(output)


def main():
    """CLI entry point: scan one or more documents passed as arguments."""
    import sys
    from app import ask_enhanced

    scanner = DDQScanner()
    os.makedirs(config.EXPORTS_DIR, exist_ok=True)

    # Files to scan can be passed on the command line, e.g.
    #   python ddq_scanner.py "sample_ddq.docx" "another_ddq.pdf"
    test_files = sys.argv[1:] or ["sample_ddq.docx"]

    for file_path in test_files:
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            continue

        try:
            results = scanner.process_document(file_path, ask_enhanced)

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = os.path.splitext(os.path.basename(file_path))[0].replace(' ', '_')
            output_path = os.path.join(config.EXPORTS_DIR, f"{filename}_{timestamp}.csv")

            scanner.export_to_csv(results, output_path)
            print("\n" + "=" * 80 + "\n")

        except Exception as e:
            print(f"Error: {str(e)}")
            import traceback
            traceback.print_exc()


if __name__ == "__main__":
    main()
