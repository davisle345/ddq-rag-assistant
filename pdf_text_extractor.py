"""
PDF / DOCX text extraction.

A lightweight, dependency-free extractor built on `pdfplumber` (PDF) and
`python-docx` (Word). It exposes a `text_extractor()` method so it can be used
interchangeably by the bulk import pipeline.

Note: the original internal version of this project optionally ran OCR and a
local vision-LLM pass for scanned documents. That path depended on proprietary
tooling and is intentionally omitted here. For scanned/image-only PDFs, add an
OCR step (e.g. `ocrmypdf` or `pytesseract`) before extraction.
"""

import os
from typing import Optional

import pdfplumber


class PdfTextExtractor:
    """Extract text (and tables) from PDF or DOCX documents."""

    def __init__(self, input_path: str, output_path: Optional[str] = None):
        self.input_path = input_path
        self.output_path = output_path

    # ------------------------------------------------------------------ #
    # Extraction
    # ------------------------------------------------------------------ #
    def extract_from_pdf(self) -> str:
        """Extract text and tables from every page of a PDF."""
        if not os.path.exists(self.input_path):
            raise FileNotFoundError(f"Input file not found: {self.input_path}")

        chunks = []
        with pdfplumber.open(self.input_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                tables = page.extract_tables()
                if tables:
                    for t_idx, table in enumerate(tables, 1):
                        chunks.append(f"[Page {page_num} - Table {t_idx}]")
                        for row in table:
                            cells = [str(c) if c is not None else "" for c in row]
                            chunks.append(" | ".join(cells))
                page_text = page.extract_text()
                if page_text:
                    chunks.append(page_text)

        return "\n".join(chunks)

    def extract_from_docx(self) -> str:
        """Extract text from paragraphs and tables of a Word document."""
        from docx import Document

        if not os.path.exists(self.input_path):
            raise FileNotFoundError(f"Input file not found: {self.input_path}")

        doc = Document(self.input_path)
        chunks = [p.text for p in doc.paragraphs if p.text.strip()]

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    chunks.append(" | ".join(cells))

        return "\n".join(chunks)

    def text_extractor(self) -> Optional[str]:
        """
        Extract text from the configured input document.

        Returns the extracted text, optionally writing it to `output_path`.
        Returns None on failure.
        """
        try:
            lower = self.input_path.lower()
            if lower.endswith(".pdf"):
                text = self.extract_from_pdf()
            elif lower.endswith((".docx", ".doc")):
                text = self.extract_from_docx()
            else:
                raise ValueError(f"Unsupported file type: {self.input_path}")

            if self.output_path:
                out = self.output_path
                if not out.endswith((".md", ".markdown", ".txt")):
                    out = f"{out}.md"
                with open(out, "w", encoding="utf-8") as f:
                    f.write(text)

            return text

        except Exception as exc:  # noqa: BLE001 - surface a readable message
            print(f"Error extracting text from {self.input_path}: {exc}")
            return None


# Backwards-compatible alias for the bulk import pipeline.
PdfTextExtractorToolkit = PdfTextExtractor


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python pdf_text_extractor.py <path-to-pdf-or-docx>")
        sys.exit(1)

    extractor = PdfTextExtractor(input_path=sys.argv[1])
    extracted = extractor.text_extractor()
    print(extracted or "(no text extracted)")
